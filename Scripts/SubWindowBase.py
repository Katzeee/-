from PySide2.QtWidgets import *
from PySide2.QtUiTools import QUiLoader
from PySide2.QtCore import *
from PySide2.QtGui import *
import os
from lib.share import SI
from lib.data import *
import time
import sys
import threading
import pandas as pd
from toolkit import *
from queue import Queue
import queue
import random
import docx
import difflib
import codecs
import hashlib




class SubWindowBase(QMdiSubWindow,threading.Thread):


    def __init__(self):
        threading.Thread.__init__(self); 
        QMdiSubWindow.__init__(self)
        self.setParent(SI.MainWin.ui.mdiArea)
        self.resize(400, 470)
        self.show()
        self.windowType = None

        #注册到子窗口管理器
        #self.windowData = SubWindowData()
        SI.SubWindowManager.RegisterSubWindow(SI.MainWin.ui.mdiArea.activeSubWindow(), self)

        # 开启线程
        self.thIsOn = True
        # 终止搜索判断
        self.stopSignal = False 



    def closeEvent(self, evnet):
        mdiId = SI.MainWin.ui.mdiArea.activeSubWindow()
        SI.SubWindowManager.RemoveSubWindow(mdiId)
        self.thIsOn = False
        print("关闭")
    
    def GetWindowType(self):
        return self.windowType
    
    def Output(self, _outFileNum):
        return




class SearchDirSubWindow(SubWindowBase):

    def __init__(self):
        
        super().__init__()
        self.pwd = os.path.abspath('.')
        self.ui = QUiLoader().load(self.pwd + '.\\Ui\\searchDirWin.ui')
        self.setWidget(self.ui)
        self.setWindowTitle("文件遍历")
        self.windowType = 'SearchDirWin'

        # 按键绑定
        self.ui.buttonMyFile.clicked.connect(self.ButtonMyFileClicked)
        self.ui.buttonClear.clicked.connect(self.ButtonClearClicked)
        self.ui.tableMyFile.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        # 参数初始化
        self.allFileInfo = pd.DataFrame()
        self.dirQueue = Queue()
        self.qLock = threading.Lock()

        self.fileNameList = []
        self.fileTypeList = []
        self.fileModifyDateList = []
        self.fileSizeList = []
        self.directory = ''

    def ButtonMyFileClicked(self):
        print("ButtonMyFileClicked")
        self.directory = QFileDialog.getExistingDirectory(dir=self.pwd)

    def ButtonClearClicked(self):
        if self.ui.buttonClear.text() == '清除缓存':
            print("ButtonClearClicked")
            self.allFileInfo = pd.DataFrame()
            self.TableMyFileShow()
            self.fileNameList = []
            self.fileTypeList = []
            self.fileModifyDateList = []
            self.fileSizeList = []
            self.ui.labelState.setText("空闲")
        else:
            self.stopSignal = True

    def FileSplit(self, _directory, _iter = 0): # 层次遍历
        try:
            _directoryList = os.listdir(_directory)
        except Exception as e:
            print(e)
        else:
            _fileNameList = []
            _fileTypeList = []
            _fileModifyDateList = []
            _fileSizeList = []

            for i in range(len(_directoryList)):
                dirName = _directoryList[i]
                _directoryList[i] = _directory + '/' +dirName
                dirAddress = _directoryList[i]
                try:
                    statInfo = os.stat(dirAddress)
                except FileNotFoundError:
                    break
                if os.path.isdir(dirAddress): 
                    _fileNameList.append(dirName)
                    _fileTypeList.append('Folder')
                    _fileModifyDateList.append(stamp2string(statInfo[8]))
                    _fileSizeList.append(statInfo[6])

                else:
                    _fileNameList.append(dirName)
                    _fileTypeList.append(dirAddress.split('.')[-1])
                    _fileModifyDateList.append(stamp2string(statInfo[8]))
                    _fileSizeList.append(statInfo[6])
                
            for i in range(_iter):
                temp = []
                while _directoryList:
                    _directory = _directoryList.pop()
                    self.ui.labelState.setText(_directory)
                    time.sleep(random.random() / 1000)
                    try:
                        dirNames = os.listdir(_directory)  # 获取目录下所有一层文件与文件夹名
                    except Exception as e:
                        print(e)
                    else:
                        for dirName in dirNames:  # 遍历每个文件
                            dirAddress = _directory + '/' + dirName  # 获取绝对路径
                            try:
                                statInfo = os.stat(dirAddress)
                            except FileNotFoundError:
                                break
                            if os.path.isdir(dirAddress): 
                                _fileNameList.append(dirName)
                                _fileTypeList.append('Folder')
                                _fileModifyDateList.append(stamp2string(statInfo[8]))
                                _fileSizeList.append(statInfo[6])
                                temp.append(dirAddress)

                            else:
                                _fileNameList.append(dirName)
                                _fileTypeList.append(dirAddress.split('.')[-1])
                                _fileModifyDateList.append(stamp2string(statInfo[8]))
                                _fileSizeList.append(statInfo[6])
                _directoryList = temp

            self.qLock.acquire()
            self.fileNameList.append(_fileNameList)
            self.fileTypeList.append(_fileTypeList)
            self.fileModifyDateList.append(_fileModifyDateList)
            self.fileSizeList.append(_fileSizeList)
            self.qLock.release()
                
            return _directoryList    



    def FileFind(self, _directoryList):
        print("thread")
        _fileNameList = []
        _fileTypeList = []
        _fileModifyDateList = []
        _fileSizeList = []
        while _directoryList and not self.stopSignal:
            _directory = _directoryList.pop()
            self.ui.labelState.setText(_directory)
            time.sleep(random.random() / 1000)
            try:
                dirNames = os.listdir(_directory)  # 获取目录下所有一层文件与文件夹名
            except Exception as e:
                print(e)
            else:
                for dirName in dirNames:  # 遍历每个文件
                    dirAddress = _directory + '/' + dirName  # 获取绝对路径
                    try:
                        statInfo = os.stat(dirAddress)
                    except FileNotFoundError:
                        break
                    '''
                    statInfo返回一个tuple，os.stat_result(st_mode=33206, st_ino=1125899906860840, st_dev=4199183826, st_nlink=1, 
                    st_uid=0, st_gid=0, st_size=80, st_atime=1616486743, st_mtime=1616485640, st_ctime=1616485640)
                    采取数字索引，size(6)代表其大小，单位为字节，mtime（8）代表其修改时间
                    '''
                    if os.path.isdir(dirAddress): 
                        _fileNameList.append(dirName)
                        _fileTypeList.append('Folder')
                        _fileModifyDateList.append(stamp2string(statInfo[8]))
                        _fileSizeList.append(statInfo[6])
                        _directoryList.append(dirAddress)

                    else:
                        _fileNameList.append(dirName)
                        _fileTypeList.append(dirAddress.split('.')[-1])
                        _fileModifyDateList.append(stamp2string(statInfo[8]))
                        _fileSizeList.append(statInfo[6])

        self.qLock.acquire()
        self.fileNameList.append(_fileNameList)
        self.fileTypeList.append(_fileTypeList)
        self.fileModifyDateList.append(_fileModifyDateList)
        self.fileSizeList.append(_fileSizeList)
        self.qLock.release()

    def TableMyFileShow(self):
        #print(self.allFileInfo)
        num = len(self.allFileInfo.index)
        self.ui.tableMyFile.setRowCount(num)
        for index in range(num):
            self.ui.tableMyFile.setItem(index, 0, QTableWidgetItem(self.fileNameList[index]))
            self.ui.tableMyFile.setItem(index, 1, QTableWidgetItem(self.fileTypeList[index]))
            self.ui.tableMyFile.setItem(index, 2, QTableWidgetItem(self.fileModifyDateList[index]))
            size = self.fileSizeList[index]
            if size == '':
                pass
            else:
                size /= 1024
                if size <= 1024:
                    self.ui.tableMyFile.setItem(index, 3, QTableWidgetItem(f"{size:.2f}KB"))
                else:
                    size /= 1024
                    self.ui.tableMyFile.setItem(index, 3, QTableWidgetItem(f"{size:.2f}MB"))

    # 导出到文件
    def Output(self, _outFileNum):
        outputList = QFileDialog.getSaveFileName(dir=f"Output{_outFileNum}.csv", filter='All Files (*) ;; csv(*.csv)', selectedFilter='csv(*.csv)')
        if outputList[0] == '': #未选择
            return
        self.allFileInfo.to_csv(outputList[0], index=None) # 输出到文件

        return outputList[0].split('/')[-1] #返回文件名


        

    def run(self):
        while self.thIsOn:
            if self.directory == '':
                time.sleep(1)
            else:
                timeStart=time.time()

                self.ui.labelState.setText("搜索中...")
                self.ui.buttonClear.setText('停止搜索')

                dirList = self.FileSplit(self.directory, 3)
                dirListLen = len(dirList)
                if dirListLen != 0:
                    th = []
                    nThread = 3
                    for i in range(nThread):
                        start = int(i /nThread * dirListLen)
                        end = int((i + 1) / nThread * dirListLen)
                        #print(start, end)
                        t = threading.Thread(target = SearchDirSubWindow.FileFind, args = (self,dirList[start : end]))
                        t.daemon = True
                        t.start()
                        th.append(t)
                    
                    for t in th:
                        t.join()
                
                # 展开数据并展示
                self.fileNameList = FlattenList(self.fileNameList)
                self.fileTypeList = FlattenList(self.fileTypeList)
                self.fileModifyDateList = FlattenList(self.fileModifyDateList)
                self.fileSizeList = FlattenList(self.fileSizeList)
                self.allFileInfo = pd.DataFrame({'fileName' : self.fileNameList, 'fileType' : self.fileTypeList, 
                                                'fileModifyDate' : self.fileModifyDateList, 'fileSize' : self.fileSizeList})
                timeEnd=time.time()
                self.ui.labelState.setText(f"搜索完成，总计{len(self.fileNameList)}个文件,总耗时{timeEnd - timeStart:.3f}s")
                #QMessageBox.information(self, "Tips", f"搜索完成，总耗时{timeEnd - timeStart:.3f}s")
                self.ui.buttonClear.setText('清除缓存')
                self.TableMyFileShow()

                # 初始化变量为下次搜索准备
                self.fileNameList = [self.fileNameList]
                self.fileTypeList = [self.fileTypeList]
                self.fileModifyDateList = [self.fileModifyDateList]
                self.fileSizeList = [self.fileSizeList]
                self.directory = ''
                self.stopSignal = False
                

class SearchNameSubWindow(SubWindowBase):
    def __init__(self):
        
        super().__init__()
        self.pwd = os.path.abspath('.')
        self.ui = QUiLoader().load(self.pwd + '.\\Ui\\searchNameWin.ui')
        self.setWidget(self.ui)
        self.setWindowTitle("模糊查找")
        self.windowType = 'SearchNameWin'

        # 按键绑定
        self.ui.buttonSearch.clicked.connect(self.ButtonSearchClicked)
        self.ui.buttonChooseDir.clicked.connect(self.ButtonChooseDirClicked)
        # self.ui.tableMyFile.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        # 右键菜单初始化
        self.RightClickMenuInit()
        

        # 参数初始化
        self.fileSearchName = ''
        self.directory = ''
        self.dirQueue = Queue()
        self.qLock = threading.Lock()

    def RightClickMenuInit(self):
        self.ui.listWidgetAllFile.setContextMenuPolicy(Qt.CustomContextMenu)
        # 必须将ContextMenuPolicy设置为Qt.CustomContextMenu
        # 否则无法使用customContextMenuRequested信号
        self.ui.Menu = QMenu(self.ui.listWidgetAllFile)
        self.ui.Menu.addAction(u'打开文件').triggered.connect(self.RightClickMenuOpenFile)
        self.ui.Menu.addAction(u'打开文件夹').triggered.connect(self.RightClickMenuOpenFilePath)
        self.ui.Menu.addAction(u'删除文件').triggered.connect(self.RightClickMenuDeleteFile)
        self.ui.listWidgetAllFile.customContextMenuRequested[QPoint].connect(self.RightClickMenuClicked)

    def RightClickMenuClicked(self, _QPoint):
        self.itemSelected = self.ui.listWidgetAllFile.itemAt(_QPoint)
        if self.itemSelected is None: 
            return
        self.ui.Menu.exec_(QCursor.pos())

    def RightClickMenuOpenFile(self):
        print("RightClickMenuOpenFile")
        text = self.itemSelected.text()
        os.startfile(text)

    def RightClickMenuOpenFilePath(self):
        text = self.itemSelected.text()
        os.startfile(text.strip(text.split('/')[-1]))

    def RightClickMenuDeleteFile(self):
        print("RightClickMenuDeleteFile")


    def ButtonSearchClicked(self):
        if self.ui.buttonSearch.text() == '开始搜索':
            self.fileSearchName = self.ui.lineEditFileKeyName.text()
            if self.fileSearchName == '' or self.directory == '':
                self.ui.labelState.setText(f"Warning！请输入文件名或路径！")
                #QMessageBox.information(self, "Warning", "请输入文件名或路径")
                return
        elif self.ui.buttonSearch.text() == '停止搜索':
            self.stopSignal = True

    def ButtonChooseDirClicked(self):
        print("ButtonChooseDirClicked")
        self.directory = QFileDialog.getExistingDirectory(dir=self.pwd)
        self.ui.lineEditFileDir.setText(self.directory)


    def FileSplit(self, _directory, _iter = 0):
        try:
            _directoryList = os.listdir(_directory)
        except IOError:
            print("IOError")
        else:
            for i in range(len(_directoryList)): # 根目录
                dirName = _directoryList[i]
                _directoryList[i] = _directory + '/' +dirName
                dirAddress = _directoryList[i]
                try:
                    statInfo = os.stat(dirAddress)
                except FileNotFoundError:
                    break
                if self.fileSearchName in dirName:
                    self.ui.listWidgetAllFile.addItem(QListWidgetItem(dirAddress))
                
            for i in range(_iter): #子级目录
                temp = []
                while _directoryList:
                    _directory = _directoryList.pop()
                    self.ui.labelState.setText(_directory)
                    try:
                        dirNames = os.listdir(_directory)  # 获取目录下所有一层文件与文件夹名
                    except IOError:
                        print(_directory)
                        print("IOError")
                    else:
                        for dirName in dirNames:  # 遍历每个文件
                            dirAddress = _directory + '/' + dirName  # 获取绝对路径
                            try:
                                statInfo = os.stat(dirAddress)
                            except FileNotFoundError:
                                break
                            if self.fileSearchName in dirName:
                                self.ui.listWidgetAllFile.addItem(QListWidgetItem(dirAddress))
                            if os.path.isdir(dirAddress): 
                                temp.append(dirAddress)
                _directoryList = temp
           
            return _directoryList    



    def FileKeyNameSearch(self, _directoryList):
        while _directoryList and not self.stopSignal:
            _directory = _directoryList.pop()
            self.ui.labelState.setText(_directory)
            time.sleep(random.random() / 1000)
            try:
                dirNames = os.listdir(_directory)  # 获取目录下所有一层文件与文件夹名
            except IOError:
                print(_directory)
                print("IOError")
            else:
                for dirName in dirNames:  # 遍历每个文件
                    dirAddress = _directory + '/' + dirName  # 获取绝对路径
                    try:
                        statInfo = os.stat(dirAddress)
                    except FileNotFoundError:
                        break
                    if self.fileSearchName in dirName:
                        self.ui.listWidgetAllFile.addItem(QListWidgetItem(dirAddress))
                    if os.path.isdir(dirAddress): 
                        _directoryList.append(dirAddress)



    
    def run(self):
        while self.thIsOn:
            if self.fileSearchName == '' or self.directory == '':
                time.sleep(1)
            else:
                self.ui.labelState.setText("搜索中...")
                self.ui.buttonSearch.setText('停止搜索')

                dirList = self.FileSplit(self.directory, 0)
                dirListLen = len(dirList)
                if dirListLen != 0:
                    th = []
                    nThread = 3
                    for i in range(nThread):
                        start = int(i /nThread * dirListLen)
                        end = int((i + 1) / nThread * dirListLen)
                        #print(start, end)
                        t = threading.Thread(target = SearchNameSubWindow.FileKeyNameSearch, args = (self,dirList[start : end]))
                        t.daemon = True
                        t.start()
                        th.append(t)
                    
                    for t in th:
                        t.join()

                self.ui.labelState.setText("搜索完成")
                self.ui.buttonSearch.setText('开始搜索')
                self.directory = ''
                self.fileSearchName = ''
                self.ui.lineEditFileDir.setText(self.directory)
                self.stopSignal = True

class SearchTypeSubWindow(SubWindowBase):
    def __init__(self):
        super().__init__()
        self.pwd = os.path.abspath('.')
        self.ui = QUiLoader().load(self.pwd + '.\\Ui\\searchTypeWin.ui')
        self.setWidget(self.ui)
        self.setWindowTitle("精确查找")
        self.windowType = 'SearchTypeWin'


        # 参数初始化
        self.fileSearchName = ''
        self.directory = ''
        self.dirQueue = Queue()
        self.qLock = threading.Lock()



class SearchDiffSubWindow(SubWindowBase):
    def __init__(self):
        super().__init__()
        self.pwd = os.path.abspath('.')
        self.ui = QUiLoader().load(self.pwd + '.\\Ui\\searchDiffWin.ui')
        #self.ui.textEditCompare.setParent(self.ui)
        self.setWidget(self.ui)
        self.setWindowTitle("文档重复比对")
        self.windowType = 'SearchDiffWin'
        
        self.ui.textEditCompare.setFontPointSize(9)
        
        # 按键绑定
        self.ui.buttonChooseFile1.clicked.connect(self.ButtonChooseFile1Clicked)
        self.ui.buttonChooseFile2.clicked.connect(self.ButtonChooseFile2Clicked)
        self.ui.buttonSearch.clicked.connect(self.ButtonSearchClicked)
        # self.ui.tableMyFile.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        # 参数初始化
        self.fileName1 = ('',)
        self.fileName2 = ('',)
        #self.directory = ''
        self.dirQueue = Queue()
        self.qLock = threading.Lock()
        self.stopSignal = True
        self.deltaHtml = ''

    def ButtonChooseFile1Clicked(self):
        self.fileName1 = QFileDialog.getOpenFileName(dir=self.pwd, filter='All Files (*) ;; doc(*.doc, *.docx)', selectedFilter='doc(*.doc, *.docx)')
        if self.fileName1[1] != 'doc(*.doc, *.docx)':
            self.ui.labelState.setText("请选择正确的文件类型!")
            self.fileName1 = ('',)
            return
        self.ui.lineEditFile1.setText(self.fileName1[0])
        self.ui.labelState.setText("空闲")


    def ButtonChooseFile2Clicked(self):
        self.fileName2 = QFileDialog.getOpenFileName(dir=self.pwd, filter='All Files (*) ;; doc(*.doc, *.docx)', selectedFilter='doc(*.doc, *.docx)')
        if self.fileName2[1] != 'doc(*.doc, *.docx)':
            self.ui.labelState.setText("请选择正确的文件类型!")
            self.fileName2 = ('',)
            return
        self.ui.lineEditFile2.setText(self.fileName2[0])
        self.ui.labelState.setText("空闲")



    def ButtonSearchClicked(self):
        if self.ui.buttonSearch.text() == '开始搜索':
            if self.fileName1[0] == '' or self.fileName2[0] == '':
                self.ui.labelState.setText("文件尚未选择,请选择文件!")
                return
            else:
                self.ui.textEditCompare.clear()
                self.stopSignal = False
                self.ui.buttonSearch.setText('停止搜索')
        else:
            self.ui.buttonSearch.setText('开始搜索')
            self.stopSignal = True


    def FileCompare(self):
        
        file1 = docx.Document(self.fileName1[0])
        file2 = docx.Document(self.fileName2[0])

        file1Para = ''
        file2Para = ''
        for para in file1.paragraphs:
            file1Para += para.text + '\n'
        for para in file2.paragraphs:
            file2Para += para.text + '\n'
        
        file1ParaList = file1Para.splitlines()
        file2ParaList = file2Para.splitlines()
        
        self.deltaHtml = difflib.HtmlDiff().make_file(file1ParaList, file2ParaList)
        deltaText = difflib.Differ().compare(file1ParaList, file2ParaList)

        for row in deltaText:
            self.ui.textEditCompare.append(row)
            time.sleep(0.005)
        # with codecs.open('diff.html','w',encoding='utf-8') as f:
        #     f.write(delta_html)
        # pass
    
    def Output(self, _outFileNum):
        outputList = QFileDialog.getSaveFileName(dir=f"Output{_outFileNum}.html", filter='All Files (*) ;; html(*.html)', selectedFilter='html(*.html)')
        if outputList[0] == '': #未选择
            return
        
        with codecs.open(outputList[0],'w',encoding='utf-8') as f:
            f.write(self.deltaHtml)
        return outputList[0].split('/')[-1] #返回文件名
    
    def run(self):
        while self.thIsOn:
            if self.stopSignal:
                time.sleep(1)
            else:
                self.ui.labelState.setText("对比中...")
                self.FileCompare()
                self.ui.labelState.setText("对比完成")
                self.ui.buttonSearch.setText('开始搜索')
                self.stopSignal = True

class SearchDupSubWindow(SubWindowBase):
    def __init__(self):
        super().__init__()
        self.pwd = os.path.abspath('.')
        self.ui = QUiLoader().load(self.pwd + '.\\Ui\\searchDupWin.ui')
        self.setWidget(self.ui)
        self.setWindowTitle("重复文件查找")
        self.windowType = 'SearchDupWin'
        

        
        # 按键绑定
        self.ui.buttonChooseDir1.clicked.connect(self.ButtonChooseDir1Clicked)
        self.ui.buttonChooseDir2.clicked.connect(self.ButtonChooseDir2Clicked)
        self.ui.buttonSearch.clicked.connect(self.ButtonSearchClicked)
        # self.ui.tableMyFile.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        # 参数初始化
        self.dirName1 = ''
        self.dirName2 = ''
        self.dirDict = {}
        self.dirQueue = Queue()
        self.qLock = threading.Lock()
        self.stopSignal = True


    def ButtonChooseDir1Clicked(self):
        self.dirName1 = QFileDialog.getExistingDirectory(dir=self.pwd)
        self.ui.lineEditFile1.setText(self.dirName1)
        self.ui.labelState.setText("空闲")


    def ButtonChooseDir2Clicked(self):
        self.dirName2 = QFileDialog.getExistingDirectory(dir=self.pwd)
        self.ui.lineEditFile2.setText(self.dirName2)
        self.ui.labelState.setText("空闲")



    def ButtonSearchClicked(self):
        if self.ui.buttonSearch.text() == '开始搜索':
            if self.dirName1 == '' or self.dirName2 == '':
                self.ui.labelState.setText("文件尚未选择,请选择文件!")
                return
            else:
                self.dirDict = {}
                #self.ui.tableFileCompare.clear()
                self.stopSignal = False
                self.ui.buttonSearch.setText('停止搜索')
        else:
            self.ui.buttonSearch.setText('开始搜索')
            self.stopSignal = True

    def FileWalk(self, _dir):
        return (os.path.join(root,fn) for root,dirs,files in os.walk(_dir) for fn in files) #迭代路径

    def GetMd5(self, _fileName, _blockSize=65536):
        hash = hashlib.md5()
        with open(_fileName, "rb") as f:
            for block in iter(lambda: f.read(_blockSize), b""):
                hash.update(block)
        return hash.hexdigest()

    def FileCompare(self):
        for file in self.FileWalk(self.dirName1):
            if self.stopSignal is False:
                file = file.replace('\\', '/')
                try:
                    md5 = self.GetMd5(file)
                except:
                    continue
                else:
                    if md5 not in self.dirDict.keys():
                        self.dirDict[md5] = [file]
                    else:
                        self.dirDict[md5].append(file)
            else:
                return
                    
        for file in self.FileWalk(self.dirName2):
            if self.stopSignal is False:
                file = file.replace('\\', '/')
                try:
                    md5 = self.GetMd5(file)
                except:
                    continue
                else:
                    if md5 not in self.dirDict.keys():
                        self.dirDict[md5] = [file]
                    else:
                        self.dirDict[md5].append(file)
            else:
                return



    def TableFileCompareShow(self):
        for key in list(self.dirDict.keys()):
            if len(self.dirDict[key]) == 1:
                self.dirDict.pop(key)
            else:
                for file in self.dirDict[key]:
                    self.ui.listFileCompare.addItem(QListWidgetItem(file))
                self.ui.listFileCompare.addItem('')
    


    def Output(self, _outFileNum):
        outputList = QFileDialog.getSaveFileName(dir=f"Output{_outFileNum}.txt", filter='All Files (*) ;; txt(*.txt)', selectedFilter='txt(*.txt)')
        if outputList[0] == '': #未选择
            return
        with open(outputList[0], 'w') as f:
            for key in self.dirDict.keys():
                for file in self.dirDict[key]:
                    print(file, file=f)
                print('', file=f)
        
        return outputList[0].split('/')[-1] #返回文件名
    
    def run(self):
        while self.thIsOn:
            if self.stopSignal:
                time.sleep(1)
            else:
                self.ui.labelState.setText("对比中...")
                self.FileCompare()
                self.TableFileCompareShow()
                self.ui.labelState.setText("对比完成")
                self.ui.buttonSearch.setText('开始搜索')

                self.stopSignal = True

